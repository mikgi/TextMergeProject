"""
DOCX exporter module for exporting merged text to Word format.
Supports UTF-8 Thai language and styled document formatting.
"""

import logging
import re
from pathlib import Path
from typing import Optional, Tuple

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt

from src.config import Config
from src.text_cleaner import clean_text_line

DOCUMENT_TITLE = ""
DOCUMENT_FONT = "TH Sarabun New"
BODY_FONT_SIZE = 20
HEADING_FONT_SIZE = 22
TITLE_FONT_SIZE = 26


class DocxExporter:
    """Handles exporting content to DOCX format."""

    def __init__(self, config: Config):
        self.config = config
        self.logger = logging.getLogger(__name__)

    def export(
        self,
        content: str,
        output_path: Path,
        font_name: Optional[str] = None,
    ) -> Tuple[bool, Optional[Path]]:
        try:
            self.logger.debug("Creating DOCX document...")

            chapter_numbers = re.findall(r"บทที่\s*(\d+)", content)
            if chapter_numbers:
                output_path = output_path.parent / f"บทที่ {chapter_numbers[0]} - บทที่ {chapter_numbers[-1]}.docx"

            doc = Document()
            doc_font = font_name or DOCUMENT_FONT

            style = doc.styles["Normal"]
            style.font.name = doc_font
            style.font.size = Pt(BODY_FONT_SIZE)
            style._element.rPr.rFonts.set(qn("w:eastAsia"), doc_font)

            if DOCUMENT_TITLE:
                title = doc.add_heading(DOCUMENT_TITLE, level=0)
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = title.runs[0]
                run.font.name = doc_font
                run.font.size = Pt(TITLE_FONT_SIZE)
                run.bold = True
                run._element.rPr.rFonts.set(qn("w:eastAsia"), doc_font)

            current_file_title = None
            for raw in content.split("\n"):
                text = clean_text_line(raw)
                if not text:
                    continue

                if text.startswith("#FILE:"):
                    current_file_title = text.replace("#FILE:", "").strip()
                    continue

                if text.startswith("---"):
                    continue

                if text.startswith("บทที่") or text.casefold().startswith("chapter"):
                    if doc.paragraphs:
                        doc.add_page_break()
                    if current_file_title and text.startswith(current_file_title):
                        current_file_title = None
                    heading = doc.add_heading(text, level=1)
                    run = heading.runs[0]
                    run.font.name = doc_font
                    run.font.size = Pt(HEADING_FONT_SIZE)
                    run.bold = True
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), doc_font)
                    heading.paragraph_format.space_before = Pt(24)
                    heading.paragraph_format.space_after = Pt(12)
                    continue

                if current_file_title:
                    if doc.paragraphs:
                        doc.add_page_break()
                    heading = doc.add_heading(current_file_title, level=1)
                    run = heading.runs[0]
                    run.font.name = doc_font
                    run.font.size = Pt(HEADING_FONT_SIZE)
                    run.bold = True
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), doc_font)
                    heading.paragraph_format.space_before = Pt(24)
                    heading.paragraph_format.space_after = Pt(12)
                    current_file_title = None

                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.name = doc_font
                run.font.size = Pt(BODY_FONT_SIZE)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), doc_font)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.first_line_indent = Pt(24)

            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            self.logger.info(f"DOCX file saved successfully: {output_path}")
            return True, output_path
        except IOError as e:
            self.logger.error(f"IO error writing DOCX file: {str(e)}", exc_info=True)
            return False, None
        except Exception as e:
            self.logger.error(f"Unexpected error exporting DOCX: {str(e)}", exc_info=True)
            return False, None
